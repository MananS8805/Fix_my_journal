import requests
import os
import docx

# Define the base URL of the API
BASE_URL = "http://localhost:8001"

# Define the path to the sample manuscript
SAMPLE_MANUSCRIPT_PATH = "sample_manuscript.docx"

# Define the journal for which to transform the manuscript
JOURNAL = "nejm"

def test_nejm_transformation():
    """
    Tests the one-click transformation for the NEJM journal.
    """
    # Create a sample manuscript file
    document = docx.Document()
    document.add_heading('This is a sample manuscript', 0)

    document.add_heading('Abstract', level=1)
    document.add_paragraph('This is the abstract.')

    document.add_heading('Introduction', level=1)
    document.add_paragraph('This is the introduction.')

    document.add_heading('Methods', level=1)
    document.add_paragraph('This is the methods section.')

    document.add_heading('Results', level=1)
    document.add_paragraph('This is the results section.')

    document.add_heading('Discussion', level=1)
    document.add_paragraph('This is the discussion section.')

    document.add_heading('Conclusions', level=1)
    document.add_paragraph('This is the conclusions section.')

    document.add_heading('References', level=1)
    document.add_paragraph('[1] A. Author, B. Author, and C. Author, “Title of the article,” Journal, vol. 1, no. 1, pp. 1–10, 2022.')

    document.save(SAMPLE_MANUSCRIPT_PATH)


    # Send a request to the transform-manuscript endpoint
    with open(SAMPLE_MANUSCRIPT_PATH, "rb") as f:
        files = {"file": (SAMPLE_MANUSCRIPT_PATH, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = requests.post(
            f"{BASE_URL}/transform-manuscript",
            files=files,
            params={"journal": JOURNAL},
        )

    # Check that the request was successful
    if response.status_code != 200:
        print(response.text)
    assert response.status_code == 200
    assert response.json()["success"] is True

    # Check that the transformed manuscript was created
    file_path = response.json()["data"]["file_path"]
    assert os.path.exists(file_path)

    print("NEJM transformation test passed successfully!")

    # Clean up the created files
    os.remove(SAMPLE_MANUSCRIPT_PATH)
    os.remove(file_path)

if __name__ == "__main__":
    test_nejm_transformation()
