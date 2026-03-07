"""
Manuscript export engine.
Supports LaTeX, DOCX (via Pandoc or python-docx fallback), and PDF.

DOCX rendering applies:
  - Per-heading-level font, size, bold, italic, ALL CAPS, spacing
  - Body text styling (font, size, line spacing)
  - Inline bold/italic via **markers** and *markers* inside runs
  - Bullet / numbered list styling with indentation
  - Text alignment (justify, center, left, right) per block
  - Multi-column layout (1 or 2 columns) via Word section properties
  - Tables rendered as Word tables with bold header row
  - Rich sections structure from new parser used directly when available
  - Falls back to Markdown-based renderer for old-format manuscripts
"""

import os
import profile
from pydoc import doc
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Alignment map ─────────────────────────────────────────────────────────────

_ALIGNMENT_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _margin_to_inches(margin_str: str) -> float:
    """Convert a margin string like '1-inch', '1.25-inch' to a float."""
    try:
        return float(margin_str.replace("-inch", "").replace("inch", "").strip())
    except Exception:
        return 1.0


def _parse_line_spacing(spacing) -> float:
    """Normalise line_spacing to a float multiplier."""
    try:
        return float(spacing)
    except Exception:
        return 1.5


def _set_columns(doc: Document, num_columns: int) -> None:
    """Set multi-column layout on the first (and only) Word section."""
    if num_columns <= 1:
        return
    sectPr = doc.sections[0]._sectPr
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_columns))
    cols.set(qn("w:space"), "720")    # 720 twips ≈ 0.5 inch gutter
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def _apply_run_style(run, font_name: str, font_size: float,
                     bold=False, italic=False, all_caps=False) -> None:
    """Apply font properties to a single Run."""
    run.font.name     = font_name
    run.font.size     = Pt(font_size)
    run.font.bold     = bold
    run.font.italic   = italic
    run.font.all_caps = all_caps


def _apply_paragraph_spacing(para, space_before_pt: float = 0,
                               line_spacing: float = 1.5) -> None:
    """Set spacing before and line spacing on a paragraph."""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(space_before_pt * 20)))   # twips
    twips = int(line_spacing * 240)   # 240=single, 360=1.5x, 480=double
    spacing.set(qn("w:line"), str(twips))
    spacing.set(qn("w:lineRule"), "auto")


def _inline_runs(para, text: str, base_font: str, base_size: float,
                 base_bold=False, base_italic=False) -> None:
    """
    Parse a text string that may contain **bold** and *italic* markers
    and add correctly styled runs to the paragraph.
    Handles mixed content like 'normal **bold** and *italic* text'.
    """
    token_pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')
    parts = token_pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            _apply_run_style(run, base_font, base_size,
                             bold=True, italic=base_italic)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            _apply_run_style(run, base_font, base_size,
                             bold=base_bold, italic=True)
        else:
            run = para.add_run(part)
            _apply_run_style(run, base_font, base_size,
                             bold=base_bold, italic=base_italic)


# ── Core styling methods ──────────────────────────────────────────────────────

def _style_body(para, profile: Dict, is_first_after_heading: bool = False) -> None:
    """Apply journal body-text style."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    font_name    = profile.get("font", "Times New Roman")
    font_size    = profile.get("font_size", 12)
    line_spacing = _parse_line_spacing(profile.get("line_spacing", 1.5))
    justify      = profile.get("justify", True)

    for run in para.runs:
        _apply_run_style(run, font_name, font_size)

    _apply_paragraph_spacing(para, space_before_pt=0, line_spacing=line_spacing)

    if justify:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # First-line indent — skip on first paragraph after heading
    indent = profile.get("first_line_indent", 0)
    if indent and not is_first_after_heading:
        para.paragraph_format.first_line_indent = Inches(indent)
    else:
        para.paragraph_format.first_line_indent = Inches(0)


def _style_heading(para, level: int, profile: Dict) -> None:
    """Apply journal heading style for the given heading level (1–3)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    heading_styles = profile.get("heading_styles", {})
    style          = heading_styles.get(level, {})

    font_name    = style.get("font",     profile.get("font", "Times New Roman"))
    font_size    = style.get("size",     profile.get("font_size", 12) + (4 - level))
    bold         = style.get("bold",     True)
    italic       = style.get("italic",   False)
    all_caps     = style.get("caps",     False)
    space_before = style.get("spacing_before", 10)
    space_after  = style.get("spacing_after",  4)
    alignment    = style.get("alignment", "left")



    for run in para.runs:
        _apply_run_style(run, font_name, font_size, bold, italic, all_caps)


    _apply_paragraph_spacing(para, space_before_pt=space_before, line_spacing=1.0)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.first_line_indent = Pt(0)


    if alignment == "center":
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _style_list_item(para, profile: Dict,
                      is_ordered: bool, index: int) -> None:
    """Prefix bullet / number and apply body style."""
    journal_list_style = profile.get("list_style", "bullet")
    if is_ordered or journal_list_style == "numbered":
        prefix = f"{index}. "
    elif journal_list_style == "dash":
        prefix = "– "
    else:
        prefix = "• "

    if para.runs:
        para.runs[0].text = prefix + para.runs[0].text
    else:
        para.add_run(prefix)

    _style_body(para, profile)





# ── Table rendering ───────────────────────────────────────────────────────────

def _add_table(doc: Document, table_data: Dict, profile: Dict) -> None:
    """Render a manuscript table as a styled Word table."""
    content = table_data.get("content", [])
    caption = table_data.get("caption", "")
    if not content:
        return

    rows = len(content)
    cols = max(len(r) for r in content) if content else 1

    font_name  = profile.get("font", "Times New Roman")
    font_size  = max(profile.get("font_size", 12) - 1, 8)
    journal    = profile.get("name", "").lower()

    # ── Table style based on journal ──────────────────────────────────────────
    # IEEE / ACM / PNAS use clean borderless "booktabs" style
    # Medical journals (NEJM, JAMA, Lancet, BMJ) use bordered grid
    # Nature / Science / Cell use minimal top+bottom border only
    booktabs_journals  = ["ieee", "acm", "pnas", "arxiv"]
    grid_journals      = ["nejm", "jama", "lancet", "bmj", "elsevier", "springer", "wiley"]
    minimal_journals   = ["nature", "science", "cell", "plos", "frontiers"]

    if any(j in journal for j in booktabs_journals):
        table_style = "booktabs"
    elif any(j in journal for j in grid_journals):
        table_style = "grid"
    else:
        table_style = "minimal"

    tbl = doc.add_table(rows=rows, cols=cols)

    # ── Apply border style ────────────────────────────────────────────────────
    def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
        """Set borders on a single cell."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")

        for side, val in [("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)]:
            el = OxmlElement(f"w:{side}")
            if val == "single":
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
            elif val == "thick":
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    "12")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
            elif val == "none":
                el.set(qn("w:val"),   "nil")
            tcBorders.append(el)

        tcPr.append(tcBorders)

    def _set_cell_shading(cell, fill: str):
        """Set background color on a cell (hex color, e.g. 'E8E8E8')."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill)
        tcPr.append(shd)

    # ── Build cells ───────────────────────────────────────────────────────────
    for r_idx, row_data in enumerate(content):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)
        is_last   = (r_idx == rows - 1)

        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= cols:
                break

            cell = row.cells[c_idx]
            cell.text = ""  # clear default

            # ── Paragraph + run ───────────────────────────────────────────────
            para = cell.paragraphs[0]
            run  = para.add_run(str(cell_text))
            run.font.name  = font_name
            run.font.size  = Pt(font_size)
            run.font.bold  = is_header

            # ── Cell alignment ────────────────────────────────────────────────
            # Headers: centered. Body: left-aligned. Numeric columns: right-aligned.
            cell_str = str(cell_text).strip()
            is_numeric = bool(re.match(r"^[\d\.\,\%\$\±\-\+]+$", cell_str))

            if is_header:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif is_numeric:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ── Cell padding ──────────────────────────────────────────────────
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side in ["top", "bottom", "left", "right"]:
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"),    "80")   # ~56 twips = ~0.04"
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            tcPr.append(tcMar)

            # ── Borders per style ─────────────────────────────────────────────
            if table_style == "grid":
                # Full grid — all cells have single borders
                _set_cell_border(cell,
                    top="single", bottom="single",
                    left="single", right="single")
                if is_header:
                    _set_cell_shading(cell, "E8E8E8")  # light grey header

            elif table_style == "booktabs":
                # Top border on header, bottom border on header and last row
                # No vertical borders at all
                if is_header:
                    _set_cell_border(cell,
                        top="thick", bottom="single",
                        left="none", right="none")
                    _set_cell_shading(cell, "F5F5F5")
                elif is_last:
                    _set_cell_border(cell,
                        top="none", bottom="thick",
                        left="none", right="none")
                else:
                    _set_cell_border(cell,
                        top="none", bottom="none",
                        left="none", right="none")

            elif table_style == "minimal":
                # Only top and bottom borders on header row, bottom on last row
                if is_header:
                    _set_cell_border(cell,
                        top="single", bottom="single",
                        left="none", right="none")
                elif is_last:
                    _set_cell_border(cell,
                        top="none", bottom="single",
                        left="none", right="none")
                else:
                    _set_cell_border(cell,
                        top="none", bottom="none",
                        left="none", right="none")

    # ── Row height ────────────────────────────────────────────────────────────
    for row in tbl.rows:
        row.height = Pt(font_size * 1.8)

    # ── Caption ───────────────────────────────────────────────────────────────
    if caption:
        cap_para = doc.add_paragraph()
        cap_run  = cap_para.add_run(f"Table: {caption}")
        cap_run.font.name   = font_name
        cap_run.font.size   = Pt(font_size - 1)
        cap_run.font.italic = True
        cap_para.alignment  = WD_ALIGN_PARAGRAPH.CENTER
        _apply_paragraph_spacing(cap_para, space_before_pt=4, line_spacing=1.0)

# ── Markdown → DOCX fallback ──────────────────────────────────────────────────

def _markdown_to_docx_fallback(markdown_text: str,
                                output_path: str,
                                profile: Dict,
                                tables: List[Dict] = None) -> str:
    """
    Convert Markdown manuscript to a styled DOCX using python-docx.
    - Title, authors, abstract → single full-width column
    - Body (Introduction onwards) → journal column layout (1 or 2)
    - Tables rendered inline at original position via [TABLE X:] placeholders.
    - Inline **bold** and *italic* preserved as Word runs.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    margin = _margin_to_inches(profile.get("margins", "1-inch"))
    for section in doc.sections:
        section.top_margin    = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin   = Inches(margin)
        section.right_margin  = Inches(margin)

    # ── Start single-column (title/abstract section) ──────────────────────────
    num_columns     = profile.get("columns", 1)
    in_header       = True   # True until we hit first real body H1
    columns_applied = False
    body_start_para_index=None
    _set_columns(doc, 1)

    font_name = profile.get("font", "Times New Roman")
    font_size = profile.get("font_size", 12)

    lines            = markdown_text.split("\n")
    list_counter     = 0
    in_list          = False
    last_was_heading = True

    
    for raw_line in lines:
        line = raw_line.rstrip()

        # ── H3 ────────────────────────────────────────────────────────────────
        if line.startswith("### "):
            in_list = False; list_counter = 0
            text = re.sub(r':\s*$', '', line[4:].strip())
            para = doc.add_heading(text, level=3)
            _style_heading(para, 3, profile)
            last_was_heading = True

        # ── H2 ────────────────────────────────────────────────────────────────
        elif line.startswith("## "):
            in_list = False; list_counter = 0
            text = re.sub(r':\s*$', '', line[3:].strip())
            para = doc.add_heading(text, level=2)
            _style_heading(para, 2, profile)
            last_was_heading = True

        # ── H1 ────────────────────────────────────────────────────────────────
        elif line.startswith("# "):
            in_list = False; list_counter = 0
            text = re.sub(r':\s*$', '', line[2:].strip())

            # Switch to multi-column on first real body section heading
            if (in_header
                    and not columns_applied
                    and num_columns > 1
                    and text.upper() not in ("ABSTRACT", "KEYWORDS", "AUTHORS")):
                print(f"COLUMN SWITCH triggered at heading: '{text}' para_index={len(doc.paragraphs)}")
                body_start_para_index= len(doc.paragraphs)
                columns_applied = True
                in_header       = False

            para = doc.add_heading(text, level=1)
            _style_heading(para, 1, profile)
            last_was_heading = True

        # ── Ordered list ──────────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", line):
            in_list = True
            list_counter += 1
            text = re.sub(r"^\d+\.\s+", "", line)
            para = doc.add_paragraph()
            para.add_run(text)
            _style_list_item(para, profile, is_ordered=True, index=list_counter)
            last_was_heading = False

        # ── Unordered list ────────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            in_list = True
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.add_run(text)
            _style_list_item(para, profile, is_ordered=False, index=0)
            last_was_heading = False

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            in_list = False; list_counter = 0
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)
            last_was_heading = False

        # ── Empty line ────────────────────────────────────────────────────────
        elif line.strip() == "":
            in_list = False; list_counter = 0
            para = doc.add_paragraph()
            _apply_paragraph_spacing(para, space_before_pt=0, line_spacing=1.0)
            # Don't reset last_was_heading on blank lines

        # ── Table placeholder — render inline at original position ────────────
        elif re.match(r'^\[TABLE \d+:', line.strip()):
            in_list   = False
            tbl_match = re.match(r'^\[TABLE (\d+):', line.strip())
            if tbl_match and tables:
                idx = int(tbl_match.group(1))
                if idx < len(tables):
                    doc.add_paragraph()
                    _add_table(doc, tables[idx], profile)
                    doc.add_paragraph()
            last_was_heading = False

        # ── Regular body paragraph ────────────────────────────────────────────
        else:
            in_list = False
            para    = doc.add_paragraph()
            _inline_runs(para, line, font_name, font_size)
            _style_body(para, profile, is_first_after_heading=last_was_heading)
            last_was_heading = False

    
    # ── Apply 2-column layout to body section ────────────────────────────
    if num_columns > 1 and body_start_para_index is not None:
        try:
            split_para = doc.paragraphs[body_start_para_index - 1] \
                        if body_start_para_index > 0 \
                        else doc.paragraphs[0]
            pPr    = split_para._p.get_or_add_pPr()
            sectPr = OxmlElement("w:sectPr")
            type_el = OxmlElement("w:type")
            type_el.set(qn("w:val"), "continuous")
            sectPr.append(type_el)
            cols_el = OxmlElement("w:cols")
            cols_el.set(qn("w:num"),        str(num_columns))
            cols_el.set(qn("w:space"),      "720")
            cols_el.set(qn("w:equalWidth"), "1")
            sectPr.append(cols_el)
            pgSz = OxmlElement("w:pgSz")
            pgSz.set(qn("w:w"), "12240")
            pgSz.set(qn("w:h"), "15840")
            sectPr.append(pgSz)
            pPr.append(sectPr)
        except Exception as e:
            print(f"Column break insertion failed: {e}")

    doc.save(output_path)
    return output_path
def _markdown_to_docx_fallback(markdown_text: str,
                                output_path: str,
                                profile: Dict,
                                tables: List[Dict] = None) -> str:
    """
    Convert Markdown manuscript to a styled DOCX using python-docx.
    - Title, authors, abstract → single full-width column
    - Body (Introduction onwards) → journal column layout (1 or 2)
    - Tables rendered inline at original position via [TABLE X:] placeholders.
    - Inline **bold** and *italic* preserved as Word runs.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    margin = _margin_to_inches(profile.get("margins", "1-inch"))
    for section in doc.sections:
        section.top_margin    = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin   = Inches(margin)
        section.right_margin  = Inches(margin)

    # ── Start single-column (title/abstract section) ──────────────────────────
    num_columns     = profile.get("columns", 1)
    in_header       = True   # True until we hit first real body H1
    columns_applied = False
    body_start_para_index=None
    _set_columns(doc, 1)

    font_name = profile.get("font", "Times New Roman")
    font_size = profile.get("font_size", 12)

    lines            = markdown_text.split("\n")
    list_counter     = 0
    in_list          = False
    last_was_heading = True

    
    for raw_line in lines:
        line = raw_line.rstrip()

        # ── H3 ────────────────────────────────────────────────────────────────
        if line.startswith("### "):
            in_list = False; list_counter = 0
            text = re.sub(r':\s*$', '', line[4:].strip())
            para = doc.add_heading(text, level=3)
            _style_heading(para, 3, profile)
            last_was_heading = True

        # ── H2 ────────────────────────────────────────────────────────────────
        elif line.startswith("## "):
            in_list = False; list_counter = 0
            text = re.sub(r':\s*$', '', line[3:].strip())
            para = doc.add_heading(text, level=2)
            _style_heading(para, 2, profile)
            last_was_heading = True

        # ── H1 ────────────────────────────────────────────────────────────────
        elif line.startswith("# "):
            in_list = False; list_counter = 0
            text = re.sub(r':\s*$', '', line[2:].strip())

            # Switch to multi-column on first real body section heading
            if (in_header
                    and not columns_applied
                    and num_columns > 1
                    and text.upper() not in ("ABSTRACT", "KEYWORDS", "AUTHORS")):
                body_start_para_index= len(doc.paragraphs)
                columns_applied = True
                in_header       = False

            para = doc.add_heading(text, level=1)
            _style_heading(para, 1, profile)
            last_was_heading = True

        # ── Ordered list ──────────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", line):
            in_list = True
            list_counter += 1
            text = re.sub(r"^\d+\.\s+", "", line)
            para = doc.add_paragraph()
            para.add_run(text)
            _style_list_item(para, profile, is_ordered=True, index=list_counter)
            last_was_heading = False

        # ── Unordered list ────────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            in_list = True
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.add_run(text)
            _style_list_item(para, profile, is_ordered=False, index=0)
            last_was_heading = False

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            in_list = False; list_counter = 0
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)
            last_was_heading = False

        # ── Empty line ────────────────────────────────────────────────────────
        elif line.strip() == "":
            in_list = False; list_counter = 0
            para = doc.add_paragraph()
            _apply_paragraph_spacing(para, space_before_pt=0, line_spacing=1.0)
            # Don't reset last_was_heading on blank lines

        # ── Table placeholder — render inline at original position ────────────
        elif re.match(r'^\[TABLE \d+:', line.strip()):
            in_list   = False
            tbl_match = re.match(r'^\[TABLE (\d+):', line.strip())
            if tbl_match and tables:
                idx = int(tbl_match.group(1))
                if idx < len(tables):
                    doc.add_paragraph()
                    _add_table(doc, tables[idx], profile)
                    doc.add_paragraph()
            last_was_heading = False

        # ── Regular body paragraph ────────────────────────────────────────────
        else:
            in_list = False
            para    = doc.add_paragraph()
            _inline_runs(para, line, font_name, font_size)
            _style_body(para, profile, is_first_after_heading=last_was_heading)
            last_was_heading = False

    
    # ── Apply 2-column layout to body section ────────────────────────────
    if num_columns > 1 and body_start_para_index is not None:
        try:
            split_para = doc.paragraphs[body_start_para_index - 1] \
                        if body_start_para_index > 0 \
                        else doc.paragraphs[0]
            pPr    = split_para._p.get_or_add_pPr()
            sectPr = OxmlElement("w:sectPr")
            type_el = OxmlElement("w:type")
            type_el.set(qn("w:val"), "continuous")
            sectPr.append(type_el)
            cols_el = OxmlElement("w:cols")
            cols_el.set(qn("w:num"),        str(num_columns))
            cols_el.set(qn("w:space"),      "720")
            cols_el.set(qn("w:equalWidth"), "1")
            sectPr.append(cols_el)
            pgSz = OxmlElement("w:pgSz")
            pgSz.set(qn("w:w"), "12240")
            pgSz.set(qn("w:h"), "15840")
            sectPr.append(pgSz)
            pPr.append(sectPr)
        except Exception as e:
            print(f"Column break insertion failed: {e}")

    doc.save(output_path)
    return output_path

# ── Pandoc path ───────────────────────────────────────────────────────────────

def _pandoc_available() -> bool:
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        return True
    except Exception:
        return False


def markdown_to_docx(markdown_text: str,
                     output_path: str,
                     journal: str,
                     profile: Dict,
                     tables: List[Dict] = None) -> str:
    """
    Convert Markdown to DOCX.
    Tries Pandoc + reference template first; falls back to python-docx.
    """
    template_dir  = Path(__file__).parent.parent / "templates"
    reference_doc = template_dir / f"{journal.lower()}_reference.docx"

    if _pandoc_available() and reference_doc.exists():
        try:
            with tempfile.NamedTemporaryFile(suffix=".md", delete=False,
                                             mode="w", encoding="utf-8") as f:
                f.write(markdown_text)
                md_path = f.name
            cmd = [
                "pandoc", md_path,
                "-o", output_path,
                f"--reference-doc={reference_doc}",
                "--citeproc",
            ]
            subprocess.run(cmd, check=True, capture_output=True)
            os.unlink(md_path)
            return output_path
        except Exception as e:
            print(f"Pandoc failed ({e}), falling back to python-docx.")

    return _markdown_to_docx_fallback(markdown_text, output_path, profile, tables)


# ── Rich sections → DOCX ─────────────────────────────────────────────────────

def _sections_to_docx(doc: Document, sections: list, profile: Dict) -> None:
    """
    Render the rich `sections` list from the new parser into `doc`.
    Handles all block types: paragraph, heading, list_item, list_item_ordered.
    Applies alignment, indentation, and inline bold/italic per block.
    """
    font_name    = profile.get("font",      "Times New Roman")
    font_size    = profile.get("font_size", 12)
    list_counter = {}   # {indent_level: current_count}

    for sec in sections:
        sec_level   = sec.get("level", 2)
        sec_heading = sec.get("heading", "")

        if sec_heading:
            para = doc.add_heading(sec_heading, level=sec_level)
            _style_heading(para, sec_level, profile)

        for block in sec.get("blocks", []):
            btype     = block.get("type", "paragraph")
            text      = block.get("text", "")
            blevel    = block.get("level")
            alignment = block.get("alignment", "left")
            indent    = block.get("indent", 0) or 0

            # Reset ordered counter when leaving a numbered list
            if btype != "list_item_ordered":
                list_counter = {}

            # Sub-heading inside a section
            if btype == "heading":
                hlevel = blevel if blevel else min(sec_level + 1, 3)
                para   = doc.add_heading(text, level=hlevel)
                _style_heading(para, hlevel, profile)

            # Bullet list item
            elif btype == "list_item":
                prefix = "• " if profile.get("list_style") != "dash" else "– "
                para   = doc.add_paragraph()
                _inline_runs(para, prefix + text, font_name, font_size)
                _style_body(para, profile)
                para.paragraph_format.left_indent = Inches(0.25 * (indent + 1))

            # Numbered list item
            elif btype == "list_item_ordered":
                key = indent
                list_counter[key] = list_counter.get(key, 0) + 1
                num  = list_counter[key]
                para = doc.add_paragraph()
                _inline_runs(para, f"{num}. {text}", font_name, font_size)
                _style_body(para, profile)
                para.paragraph_format.left_indent = Inches(0.25 * (indent + 1))

            # Regular paragraph
            else:
                para = doc.add_paragraph()
                _inline_runs(para, text, font_name, font_size)
                _style_body(para, profile)
                para.alignment = _ALIGNMENT_MAP.get(
                    (alignment or "justify").lower(),
                    WD_ALIGN_PARAGRAPH.JUSTIFY
                )
                if indent:
                    para.paragraph_format.left_indent = Inches(0.25 * indent)


def _rich_to_docx(manuscript: Dict,
                  output_path: str,
                  profile: Dict,
                  tables: list = None) -> str:
    """
    Build a fully styled DOCX from the rich 'sections' structure
    produced by the new parser. Falls back to Markdown renderer if
    'sections' is absent (old-format manuscripts).
    """
    sections = manuscript.get("sections", [])
    if not sections:
        body = manuscript.get("body", "")
        return _markdown_to_docx_fallback(body, output_path, profile, tables)

    doc    = Document()
    margin = _margin_to_inches(profile.get("margins", "1-inch"))
    for section in doc.sections:
        section.top_margin    = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin   = Inches(margin)
        section.right_margin  = Inches(margin)

    _set_columns(doc, profile.get("columns", 1))

    font_name = profile.get("font",      "Times New Roman")
    font_size = profile.get("font_size", 12)

    # Title
    title = re.sub(r"^#+\s*", "",
                   manuscript.get("metadata", {}).get("title", ""))
    if title:
        t_para = doc.add_heading(title, level=0)
        t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size + 4)
            run.font.bold = True

    # Abstract
    abstract = manuscript.get("abstract", "")
    if abstract:
        a_head = doc.add_heading("Abstract", level=1)
        _style_heading(a_head, 1, profile)
        a_para = doc.add_paragraph()
        _inline_runs(a_para, abstract, font_name, font_size, base_italic=True)
        _style_body(a_para, profile)
        a_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Body sections
    _sections_to_docx(doc, sections, profile)

    # Tables
    if tables:
        for table_data in tables:
            doc.add_paragraph()
            _add_table(doc, table_data, profile)

    # References
    refs = manuscript.get("references", "")
    if refs:
        r_head = doc.add_heading("References", level=1)
        _style_heading(r_head, 1, profile)
        ref_size = max(font_size - 1, 8)
        for line in refs.split("\n"):
            line = line.strip()
            if line:
                r_para = doc.add_paragraph()
                _inline_runs(r_para, line, font_name, ref_size)
                _apply_paragraph_spacing(r_para, space_before_pt=2,
                                          line_spacing=1.2)
                r_para.paragraph_format.left_indent       = Inches(0.25)
                r_para.paragraph_format.first_line_indent = Inches(-0.25)

    doc.save(output_path)
    return output_path


# ── LaTeX export ──────────────────────────────────────────────────────────────

def _generate_latex(manuscript: Dict, journal: str, profile: Dict) -> str:
    """Generate a LaTeX string from the formatted manuscript dict."""
    title    = re.sub(r"^#+\s*", "",
                      manuscript.get("metadata", {}).get("title", "Manuscript"))
    abstract = manuscript.get("abstract", "")
    body     = manuscript.get("body",     "")
    refs     = manuscript.get("references", "")

    font     = profile.get("font",      "Times New Roman")
    fontsize = profile.get("font_size", 12)
    cols     = profile.get("columns",   1)

    font_pkg = ""
    if "Times" in font:
        font_pkg = "\\usepackage{times}"
    elif "Arial" in font or "Helvetica" in font:
        font_pkg = "\\usepackage{helvet}\n\\renewcommand{\\familydefault}{\\sfdefault}"

    col_cmd = "\\twocolumn" if cols == 2 else ""

    def md_to_latex_sections(text: str) -> str:
        text = re.sub(r"^### (.+)$", r"\\subsubsection{\1}", text, flags=re.MULTILINE)
        text = re.sub(r"^## (.+)$",  r"\\subsection{\1}",    text, flags=re.MULTILINE)
        text = re.sub(r"^# (.+)$",   r"\\section{\1}",       text, flags=re.MULTILINE)
        # Inline bold/italic → LaTeX
        text = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", text)
        text = re.sub(r"\*(.+?)\*",     r"\\textit{\1}", text)
        # Bullet lists
        lines = text.split("\n")
        out, in_list = [], False
        for line in lines:
            if line.startswith("- ") or line.startswith("* "):
                if not in_list:
                    out.append("\\begin{itemize}")
                    in_list = True
                out.append(f"  \\item {line[2:]}")
            else:
                if in_list:
                    out.append("\\end{itemize}")
                    in_list = False
                out.append(line)
        if in_list:
            out.append("\\end{itemize}")
        return "\n".join(out)

    body_latex = md_to_latex_sections(body)

    latex = f"""\\documentclass[{fontsize}pt]{{article}}
{font_pkg}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{setspace}}
\\usepackage{{booktabs}}
\\title{{{title}}}
\\date{{}}
\\begin{{document}}
\\maketitle
{col_cmd}
\\begin{{abstract}}
{abstract}
\\end{{abstract}}

{body_latex}

\\section*{{References}}
{refs}
\\end{{document}}
"""
    return latex


# ── ManuscriptExporter ────────────────────────────────────────────────────────

class ManuscriptExporter:
    """
    Exports formatted manuscripts to LaTeX, DOCX, and PDF.
    Reads journal profiles for all styling decisions.
    Uses rich sections structure from new parser when available.
    """

    def __init__(self, exports_dir: str = None):
        self.exports_dir = Path(exports_dir) if exports_dir else Path("exports")
        self.exports_dir.mkdir(parents=True, exist_ok=True)

    # ── Public API ────────────────────────────────────────────────────────────

    def export_to_latex(self,
                        manuscript: Dict,
                        journal: str,
                        changelog=None,
                        profile: Dict = None) -> str:
        profile   = profile or self._load_profile(journal)
        latex_str = _generate_latex(manuscript, journal, profile)
        out_path  = str(self.exports_dir / f"{journal.lower()}_manuscript.tex")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(latex_str)
        print(f"LaTeX exported to {out_path}")
        return out_path

    def export_to_docx(self,
                       manuscript: Dict,
                       journal: str,
                       changelog=None,
                       profile: Dict = None) -> str:
        profile  = profile or self._load_profile(journal)
        tables   = manuscript.get("tables", [])
        out_path = str(self.exports_dir / f"{journal.lower()}_manuscript.docx")

        if manuscript.get("sections"):
            # New parser output — use rich renderer
            _rich_to_docx(manuscript, out_path, profile, tables)
        else:
            # Old-format manuscript — use Markdown renderer
            md_text = self._build_markdown(manuscript)
            markdown_to_docx(md_text, out_path, journal, profile, tables)

        print(f"DOCX exported to {out_path}")
        return out_path

    def export_to_pdf(self,
                      manuscript: Dict,
                      journal: str,
                      changelog=None,
                      profile: Dict = None) -> str:
        """Export to PDF via LaTeX → pdflatex compilation."""
        profile  = profile or self._load_profile(journal)
        tex_path = self.export_to_latex(manuscript, journal, changelog, profile)

        try:
            result = subprocess.run(
                ["pdflatex", "-interaction=nonstopmode",
                 "-output-directory", str(self.exports_dir), tex_path],
                capture_output=True, text=True, timeout=60,
            )
            pdf_path = tex_path.replace(".tex", ".pdf")
            if Path(pdf_path).exists():
                print(f"PDF exported to {pdf_path}")
                return pdf_path
            raise RuntimeError(
                f"pdflatex did not produce a PDF.\n{result.stderr[:500]}"
            )
        except FileNotFoundError:
            raise RuntimeError("pdflatex not found. Install TeX Live or MiKTeX.")

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _load_profile(self, journal: str) -> Dict:
        try:
            from core.journal_profiles import get_journal_profile
            return get_journal_profile(journal)
        except ImportError:
            return {}

    def _build_markdown(self, manuscript: Dict) -> str:
        title    = re.sub(r"^#+\s*", "",
                          manuscript.get("metadata", {}).get("title", "Manuscript"))
        abstract = manuscript.get("abstract", "")
        body     = manuscript.get("body",     "")
        refs     = manuscript.get("references", "")

        md  = f"# {title}\n\n"
        if abstract:
            md += f"## Abstract\n\n{abstract.strip()}\n\n"
        md += body.strip() + "\n\n"
        if refs:
            md += f"## References\n\n{refs.strip()}\n"
        return md