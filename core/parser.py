"""
Manuscript parser using Groq (Llama 3.3 70B) + PyMuPDF for local text extraction.
Uses direct httpx calls instead of groq SDK to bypass uvicorn worker network issues.
Supports large documents via chunked processing.

Now extracts rich formatting metadata:
  - Heading levels (1/2/3)
  - Inline bold/italic via Markdown markers
  - List detection (bullet / numbered)
  - Text alignment hints
  - Indentation level
"""

print("PARSER LOADED: GROQ VERSION WITH RICH FORMATTING")

import os
import json
import httpx
import time
from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
from pathlib import Path

from sympy import re


# ── Pydantic schema ────────────────────────────────────────────────────────────

class ContentBlock(BaseModel):
    """
    A single content block inside a section.
    type: "paragraph" | "heading" | "list_item" | "list_item_ordered"
    """
    type: str = Field(
        description=(
            "Block type: 'paragraph', 'heading', "
            "'list_item' (bullet), 'list_item_ordered' (numbered)"
        )
    )
    text: str = Field(
        description=(
            "Text content. Use **word** for bold, *word* for italic. "
            "Preserve original emphasis."
        )
    )
    level: Optional[int] = Field(
        None,
        description=(
            "For 'heading' blocks: 1=major, 2=subsection, 3=sub-subsection. "
            "For 'list_item_ordered': the item number (1, 2, 3...)."
        )
    )
    alignment: Optional[str] = Field(
        None,
        description="Text alignment: 'left', 'center', 'right', 'justify'. Default left."
    )
    indent: Optional[int] = Field(
        None,
        description="Indentation level: 0=none, 1=one level in, 2=two levels. Default 0."
    )


class Section(BaseModel):
    heading: str = Field(description="The section heading text (plain, no # symbols).")
    level: int = Field(
        description=(
            "Heading level: 1=top-level section (Introduction, Methods, Results...), "
            "2=subsection (2.1, Literature Review...), "
            "3=sub-subsection (2.1.1, detailed sub-topic...)."
        )
    )
    blocks: List[ContentBlock] = Field(
        description=(
            "Ordered list of content blocks in this section. "
            "Break the section content into logical blocks: "
            "paragraphs, nested headings, bullet lists, numbered lists."
        )
    )
    font_size: Optional[int] = Field(None, description="Dominant font size of body text in this section.")
    font_style: Optional[str] = Field(None, description="Dominant font style: 'regular', 'bold', 'italic'.")


class Table(BaseModel):
    caption: str = Field(description="The caption of the table.")
    content: List[List[str]] = Field(description="The table content as a 2D JSON array.")
    page: int = Field(description="The page number where the table is found.")

def _inject_tables_into_sections(raw_text: str, sections: list, tables: list) -> list:
    import re, copy
    if not tables:
        return sections
    raw_lines = raw_text.split("\n")
    anchors = []
    for i, raw_line in enumerate(raw_lines):
        m = re.match(r'^\[TABLE (\d+):', raw_line.strip())
        if m:
            idx = int(m.group(1))
            anchor = ""
            for j in range(i - 1, max(i - 10, -1), -1):
                stripped = raw_lines[j].strip()
                if stripped and not re.match(r'^\[TABLE', stripped):
                    anchor = stripped[:60].lower()
                    break
            anchors.append((idx, anchor))
    if not anchors:
        if sections:
            for idx in range(len(tables)):
                sections[-1]["blocks"].append({"type": "table", "table_index": idx})
        return sections
    sections = copy.deepcopy(sections)
    injected = set()
    for tbl_idx, anchor in anchors:
        if tbl_idx >= len(tables):
            continue
        if not anchor:
            sections[-1]["blocks"].append({"type": "table", "table_index": tbl_idx})
            injected.add(tbl_idx)
            continue
        best_sec_i = best_block_i = None
        best_score = 0
        for si, sec in enumerate(sections):
            for bi, block in enumerate(sec.get("blocks", [])):
                text = block.get("text", "").lower()
                if not text:
                    continue
                score = len(set(anchor.split()) & set(text[:60].split()))
                if score > best_score:
                    best_score = score; best_sec_i = si; best_block_i = bi
        if best_sec_i is not None and best_score >= 2:
            sections[best_sec_i]["blocks"].insert(best_block_i + 1,
                {"type": "table", "table_index": tbl_idx})
        else:
            sections[-1]["blocks"].append({"type": "table", "table_index": tbl_idx})
        injected.add(tbl_idx)
    for idx in range(len(tables)):
        if idx not in injected:
            sections[-1]["blocks"].append({"type": "table", "table_index": idx})
    return sections


def _dedup_sections(sections: list) -> list:
    import re
    def _norm(t):
        t = re.sub(r'[^a-z\s]', '', t.lower())
        return re.sub(r'\s+', ' ', t).strip()
    def _first_text(sec):
        for b in sec.get("blocks", []):
            t = b.get("text", "").strip()
            if t: return _norm(t)[:80]
        return ""
    def _len(sec):
        return sum(len(b.get("text","")) for b in sec.get("blocks",[]))
    # Pass 1: merge consecutive same-heading sections
    merged = []
    for sec in sections:
        if merged and _norm(merged[-1]["heading"]) == _norm(sec["heading"]):
            existing = {_norm(b.get("text","")) for b in merged[-1]["blocks"]}
            for b in sec.get("blocks", []):
                if _norm(b.get("text","")) not in existing:
                    merged[-1]["blocks"].append(b)
                    existing.add(_norm(b.get("text","")))
        else:
            merged.append(dict(sec))
    # Pass 2: drop non-consecutive duplicates
    seen, result = [], []
    for sec in merged:
        nh, ft = _norm(sec["heading"]), _first_text(sec)
        dup = False
        for i, (sh, sf) in enumerate(seen):
            if sh == nh:
                if ft and sf and (ft[:40] in sf or sf[:40] in ft):
                    dup = True
                    if _len(sec) > _len(result[i]):
                        result[i] = sec; seen[i] = (nh, ft)
                    break
                elif not ft or not sf:
                    dup = True; break
        if not dup:
            seen.append((nh, ft)); result.append(sec)
    removed = len(sections) - len(result)
    if removed:
        print(f"DEDUP: {len(sections)} → {len(result)} sections ({removed} removed)")
    return result



class Manuscript(BaseModel):
    title: str = Field(description="The primary title of the manuscript.")
    abstract: str = Field(description="The complete abstract of the manuscript.")
    authors: List[str] = Field(default=[], description="Full names of all authors.")
    affiliations: List[str] = Field(default=[], description="Author affiliations/institutions.")
    sections: List[Section] = Field(description="A list of all sections in the document.")
    tables: List[Table] = Field(default=[],description="A list of all tables extracted from the document.")
    citations: List[str] = Field(default=[],description="A list of all bibliographic citations or references.")
    keywords: List[str] = Field(default=[],description="A list of keywords associated with the manuscript.")


# ── Parser ─────────────────────────────────────────────────────────────────────

class ManuscriptParser:

    GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
    MODEL       = "llama-3.1-8b-instant"   # faster model
    CHUNK_SIZE  = 10000                     # smaller chunks
    OVERLAP     = 200

    def __init__(self, api_key=None):
        self.api_key = (api_key or os.getenv("GROQ_API_KEY", "")).strip()
        if not self.api_key:
            raise ValueError("GROQ_API_KEY environment variable not set.")

    # ── Text extraction ────────────────────────────────────────────────────────

    def _extract_text(self, source):
        source = str(source)
        if source.lower().endswith(".pdf"):
            import fitz
            doc = fitz.open(source)
            pages_text = [page.get_text() for page in doc]
            doc.close()
            return "\n\n".join(pages_text), [], [], []

        elif source.lower().endswith((".docx", ".doc")):
            from docx import Document as DocxDocument
            from docx.text.paragraph import Paragraph
            from docx.table import Table

            doc = DocxDocument(source)
            parts = []
            tables = []
            table_index = 0

            # ── Extract authors & affiliations directly from Python ────────────
            # Scan the first ~10 paragraphs (before Abstract) for author/affil
            # patterns. This is far more reliable than asking the LLM to find them
            # in a raw text dump.
            extracted_authors      = []
            extracted_affiliations = []
            _HEADING_STYLES = {"heading 1", "heading 2", "heading 3",
                               "title", "main heading"}
            _STOP_STYLES    = {"main heading", "content"}  # abstract/body start
            _AFFIL_KEYWORDS = {"university", "department", "institute",
                               "college", "school", "faculty", "lab",
                               "centre", "center", "research", "india",
                               "usa", "uk", "email", "@"}

            for pi, para in enumerate(doc.paragraphs[:12]):
                text  = para.text.strip()
                style = para.style.name.lower()
                if not text:
                    continue
                # Stop when we hit the abstract or body
                if style in _STOP_STYLES and pi > 1:
                    break
                if style in _HEADING_STYLES:
                    continue   # skip title heading itself
                # Affiliation heuristic: contains institution keywords
                tl = text.lower()
                if any(kw in tl for kw in _AFFIL_KEYWORDS):
                    extracted_affiliations.append(text)
                # Author heuristic: short line (< 60 chars), no institution
                # keywords, appears before affiliation lines
                elif len(text) < 80 and not extracted_affiliations:
                    extracted_authors.append(text.rstrip())

            for block in doc.element.body:
                tag = block.tag.split('}')[-1]

                if tag == 'p':
                    para = Paragraph(block, doc)
                    if para.text.strip():
                        parts.append(para.text)

                elif tag == 'tbl':
                    table = Table(block, doc)
                    # Extract table content directly in Python
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip().replace('\n', ' ')
                                for cell in row.cells]
                        rows.append(cells)

                    # Use first row as caption hint
                    caption = f"Table {table_index + 1}"
                    if rows:
                        caption = f"Table {table_index + 1}: {rows[0][0][:50]}"

                    tables.append({
                        "caption": caption,
                        "content": rows,
                        "page": 1
                    })

                    # Put a placeholder in text so body position is noted
                    parts.append(f"[TABLE {table_index}: {caption}]")
                    table_index += 1

            print(f"EXTRACTED AUTHORS: {extracted_authors}")
            print(f"EXTRACTED AFFILIATIONS: {extracted_affiliations}")
            return "\n\n".join(parts), tables, extracted_authors, extracted_affiliations

        else:
            raise ValueError(f"Unsupported file type: {Path(source).suffix}")
        return "", [], [], []  # unreachable but satisfies type checker
    # ── Groq API call ──────────────────────────────────────────────────────────

    def _call_groq(self, prompt):
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": self.MODEL,
            "messages": [{"role": "user", "content": prompt}],
            "response_format": {"type": "json_object"},
            "temperature": 0.1,
            "max_tokens": 2000,
        }
        with httpx.Client(timeout=httpx.Timeout(120.0), verify=False, follow_redirects=True) as client:
            print("Making direct httpx POST to Groq API...")
            response = client.post(self.GROQ_API_URL, headers=headers, json=payload)
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]

    # ── Main entry point ───────────────────────────────────────────────────────

    def parse(self, source):
        print(f"Starting parsing for: {source}")
        try:
            raw_text, extracted_tables, py_authors, py_affiliations = self._extract_text(source)
            print(f"Extracted {len(raw_text)} characters, {len(extracted_tables)} tables from document.")
        except Exception as e:
            print(f"Text extraction failed: {e}")
            return {"error": f"Text extraction failed: {str(e)}"}

        if not raw_text.strip():
            return {"error": "Document appears to be empty or could not be read."}

        if len(raw_text) <= self.CHUNK_SIZE:
            result = self._parse_chunk(raw_text, is_first_chunk=True, is_last_chunk=True)
        else:
            result = self._parse_large_document(raw_text)

        # Python-extracted authors/affiliations are authoritative — override LLM
        if py_authors:
            result.setdefault("metadata", {})["authors"] = py_authors
        if py_affiliations:
            result.setdefault("metadata", {})["affiliations"] = py_affiliations

        if extracted_tables:
            sections = result.get("sections", [])
            if sections:
                result["sections"] = _inject_tables_into_sections(
                    raw_text, sections, extracted_tables
                )
            result["body"] = self._sections_to_markdown(result["sections"])

        # Python-extracted tables are always authoritative
        result["tables"] = extracted_tables
        print(f"TABLES FOUND: {len(extracted_tables)}")
        return result
    # ── Chunked processing ─────────────────────────────────────────────────────

    def _parse_large_document(self, raw_text):

        import re

    # Strip table placeholders before chunking — tables already extracted by Python.
    # This prevents large table text from bloating chunks and causing 400 errors.
        clean_text = re.sub(r'\[TABLE \d+:[^\]]*\]', '', raw_text)
        clean_text = re.sub(r'\n{3,}', '\n\n', clean_text).strip()
        chunks = []
        start = 0
        while start < len(clean_text):
            end = min(start + self.CHUNK_SIZE, len(clean_text))
            chunks.append(clean_text[start:end])
            if end == len(clean_text):
                break
            start += self.CHUNK_SIZE - self.OVERLAP

        print(f"Document split into {len(chunks)} chunks.")

        title        = ""
        abstract     = ""
        authors      = []
        affiliations = []
        keywords     = []
        all_sections = []
        all_tables   = []
        references   = ""

        for i, chunk in enumerate(chunks):
            is_first = (i == 0)
            is_last  = (i == len(chunks) - 1)
            print(f"Processing chunk {i+1}/{len(chunks)}...")
            print(f"Chunk sizes: {[len(c) for c in chunks]}")

            if i > 0:
                print("Waiting 20s to respect rate limits...")
                time.sleep(20)

            for attempt in range(3):
                try:
                    result = self._parse_chunk(chunk,
                                               is_first_chunk=is_first,
                                               is_last_chunk=is_last)
                    if is_first:
                        title        = result.get("metadata", {}).get("title", "")
                        abstract     = result.get("abstract", "")
                        authors      = result.get("metadata", {}).get("authors", [])
                        affiliations = result.get("metadata", {}).get("affiliations", [])
                        keywords     = result.get("metadata", {}).get("keywords", [])

                    all_sections.extend(result.get("sections", []))
                    all_tables.extend(result.get("tables", []))

                    if is_last:
                        references = result.get("references", "")
                    break

                except Exception as e:
                    if "429" in str(e) and attempt < 2:
                        wait = 30 * (attempt + 1)
                        print(f"Rate limited. Retrying in {wait}s... (attempt {attempt+2}/3)")
                        time.sleep(wait)
                    else:
                        print(f"Chunk {i+1} failed: {e} - skipping.")
                        break

        return {
            "metadata":   {"title": title, "authors": authors, "affiliations": affiliations, "keywords": keywords},
            "abstract":   abstract,
            "body":       self._sections_to_markdown(all_sections),
            "sections":   _dedup_sections(all_sections),
            "references": references,
            "figures":    [],
            "tables":     all_tables,
        }

    # ── Single chunk parser ────────────────────────────────────────────────────

    def _parse_chunk(self, text, is_first_chunk=False, is_last_chunk=False):
        if is_first_chunk and is_last_chunk:
            context_hint = "This is the complete document."
        elif is_first_chunk:
            context_hint = "This is the BEGINNING of the document. Extract title, abstract, keywords, and early sections."
        elif is_last_chunk:
            context_hint = "This is the END of the document. Extract remaining sections and the references/bibliography."
        else:
            context_hint = "This is a MIDDLE section of the document. Extract all sections you can find."

        prompt = f"""You are an expert academic document parser. {context_hint}

Analyze the manuscript text below and return a richly structured JSON object.
Return ONLY valid JSON — no explanation, no markdown fences, no backticks.

=== JSON SCHEMA ===
{{
  "title": "paper title — empty string if not in this chunk",
  "abstract": "full abstract text — empty string if not in this chunk",
  "keywords": ["keyword1", "keyword2"],
  "sections": [
    {{
      "heading": "Section heading text — plain text, no # symbols",
      "level": 1,
      "font_size": null,
      "font_style": "regular",
      "blocks": [
        {{
          "type": "paragraph",
          "text": "Paragraph text. Use **word** for bold, *word* for italic wherever emphasis exists in the original.",
          "alignment": "justify",
          "indent": 0
        }},
        {{
          "type": "list_item",
          "text": "A bullet point item",
          "indent": 1
        }},
        {{
          "type": "list_item_ordered",
          "text": "A numbered list item",
          "level": 1,
          "indent": 1
        }},
        {{
          "type": "heading",
          "text": "A sub-heading that appears inside section body",
          "level": 3,
          "alignment": "left"
        }}
      ]
    }}
  ],
  "tables": [
    {{
      "caption": "table caption or description",
      "content": [["header1", "header2"], ["row1col1", "row1col2"]],
      "page": 1
    }}
  ],
  "citations": ["each reference entry as a string — only if this chunk contains a reference list"]
}}

=== FORMATTING RULES ===

HEADING LEVELS — assign based on document hierarchy:
  level 1 → major top-level sections: Introduction, Related Work, Methods,
             Results, Discussion, Conclusion, References, Abstract
  level 2 → numbered subsections (2.1, 3.2) or named subsections
             (Literature Review, Data Collection, Experimental Setup)
  level 3 → sub-subsections (2.1.1) or deeply nested topics

BLOCK TYPES — choose the right type for each piece of content:
  "paragraph"          → regular body text
  "heading"            → a sub-heading that appears INSIDE a section's content area
  "list_item"          → unordered/bullet list item
                         (original uses •, -, *, or similar symbols)
  "list_item_ordered"  → numbered list item
                         (original uses 1. 2. 3. or (a) (b) (c) or i. ii. iii.)

INLINE FORMATTING — apply inside the "text" field only:
  **word**  → bold   (key terms, important phrases, defined terms, labels)
  *word*    → italic (latin terms, titles, variable names, foreign words,
                      light emphasis)
  Do NOT add emphasis that wasn't present in the original document.
  Regular body text stays as plain text.

ALIGNMENT — choose based on how the text sits on the page:
  "justify" → standard body paragraphs (most common in academic papers)
  "center"  → figure/table captions, display equations, section dividers
  "left"    → headings, lists, code, most structured content
  "right"   → rare, only if clearly right-aligned in original

INDENT — indentation depth:
  0 → normal body text at page margin
  1 → one level indented (list items, block quotes, sub-content)
  2 → two levels (nested list items, deeply indented content)

IMPORTANT RULES:
- Preserve ALL content — never summarize, shorten, or truncate any section
- Every subsection (2.1, 2.2, 3.1 etc.) must be its own section entry with its own blocks
- If a paragraph introduces a list, emit it as a paragraph block followed by
  individual list_item or list_item_ordered blocks — do not merge them into one paragraph
- Tables belong in the "tables" array, not inside blocks
- font_size and font_style are optional — only set them if clearly inferable
  (e.g. abstract text is often 10pt, body is 12pt)

=== MANUSCRIPT TEXT ===
{text}"""

        raw_json = self._call_groq(prompt)
        result   = json.loads(raw_json)
        print(f"CHUNK TABLES: {len(result.get('tables', []))}")
        print(f"CHUNK TABLE SAMPLE: {result.get('tables', [{}])[0] if result.get('tables') else 'NONE'}")
        manuscript = Manuscript(**result)
        return self._to_legacy_format(manuscript)

    # ── Format conversion ──────────────────────────────────────────────────────

    def _to_legacy_format(self, manuscript: Manuscript) -> Dict[str, Any]:
        """
        Converts Manuscript → pipeline dict.

        Returns both:
          'body'     — Markdown string (backward compat, LaTeX export)
          'sections' — rich structured list (export.py uses this directly for DOCX)
        """
        rich_sections = [s.model_dump() for s in manuscript.sections]

        return {
            "metadata": {
                "title":        manuscript.title,
                "authors":      manuscript.authors,
                "affiliations": manuscript.affiliations,
                "keywords":     manuscript.keywords,
            },
            "abstract":   manuscript.abstract,
            "body":       self._sections_to_markdown(rich_sections),
            "sections":   rich_sections,
            "references": "\n".join(manuscript.citations),
            "figures":    [],
            "tables":     [t.model_dump() for t in manuscript.tables],
        }

    def _sections_to_markdown(self, sections: List[Dict]) -> str:
        """
        Convert rich section list → Markdown string.
        Used as the 'body' fallback and for LaTeX export.
        Preserves heading hierarchy, lists, and inline bold/italic.
        """
        lines = []

        for sec in sections:
            level   = sec.get("level", 2)
            heading = sec.get("heading", "")
            prefix  = "#" * max(1, min(level, 6))

            lines.append(f"{prefix} {heading}")
            lines.append("")

            for block in sec.get("blocks", []):
                btype  = block.get("type", "paragraph")
                text   = block.get("text", "")
                indent = block.get("indent", 0) or 0
                pad    = "  " * indent

                if btype == "paragraph":
                    lines.append(pad + text)
                    lines.append("")

                elif btype == "heading":
                    blevel = block.get("level", level + 1)
                    blevel = max(1, min(blevel, 6))
                    lines.append(f"{'#' * blevel} {text}")
                    lines.append("")

                elif btype == "list_item":
                    lines.append(f"{pad}- {text}")

                elif btype == "list_item_ordered":
                    item_num = block.get("level", 1)
                    lines.append(f"{pad}{item_num}. {text}")

                elif btype == "table":
                    tbl_idx = block.get("table_index", 0)
                    lines.append(f"[TABLE {tbl_idx}: Table {tbl_idx+1}]")
                    lines.append("")

                else:
                    lines.append(pad + text)
                    lines.append("")

            lines.append("")

        return "\n".join(lines).strip()