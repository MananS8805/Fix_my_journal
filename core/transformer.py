"""
Pandoc transformer: Markdown -> DOCX using a reference doc template.
Falls back to python-docx if pandoc is not installed or template is missing.
"""

import subprocess
import tempfile
import re
from pathlib import Path


def _markdown_to_docx_fallback(markdown_text: str, output_path: Path, profile: dict = None) -> Path:
    """
    Fallback DOCX generator using python-docx when pandoc is unavailable.
    Applies journal-specific font, size, spacing and margins from profile.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    profile = profile or {}
    margins = profile.get("margins", {})

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(margins.get("top",    1.0))
    section.bottom_margin = Inches(margins.get("bottom", 1.0))
    section.left_margin   = Inches(margins.get("left",   1.25))
    section.right_margin  = Inches(margins.get("right",  1.25))

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = profile.get("font", "Times New Roman")
    style.font.size = Pt(profile.get("font_size", 12))

    line_spacing = profile.get("line_spacing", 1.5)

    # ── Parse and render Markdown ─────────────────────────────────────────────
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading levels  # / ## / ### etc.
        h_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if h_match:
            level = len(h_match.group(1))
            text  = h_match.group(2).strip()
            heading = doc.add_heading(text, level=min(level, 4))
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            i += 1
            continue

        # Blank line → skip
        if not line.strip():
            i += 1
            continue

        # Normal paragraph — collect consecutive non-blank, non-heading lines
        para_lines = []
        while (
            i < len(lines)
            and lines[i].strip()
            and not re.match(r"^#{1,6}\s", lines[i])
        ):
            para_lines.append(lines[i].strip())
            i += 1

        if para_lines:
            p = doc.add_paragraph(" ".join(para_lines))
            fmt = p.paragraph_format
            fmt.line_spacing  = line_spacing
            fmt.space_before  = Pt(6)
            fmt.space_after   = Pt(6)
            fmt.first_line_indent = Inches(0.25)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def markdown_to_docx(
    markdown_text: str,
    output_path: Path,
    reference_doc: Path,
    use_citeproc: bool = True,
    journal_profile: dict = None,
) -> Path:
    """
    Convert Markdown text to DOCX.

    Tries pandoc first (with reference_doc template).
    Falls back to python-docx if:
      - pandoc is not installed
      - the reference_doc template does not exist
      - pandoc exits with an error

    journal_profile is passed to the fallback so it can apply
    the correct font, size, spacing and margins for the target journal.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # ── Check pandoc ──────────────────────────────────────────────────────────
    try:
        subprocess.run(
            ["pandoc", "--version"],
            capture_output=True,
            check=True,
            timeout=10,
        )
        pandoc_available = True
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        pandoc_available = False

    # ── Check template ────────────────────────────────────────────────────────
    template_exists = reference_doc is not None and Path(reference_doc).exists()

    if pandoc_available and template_exists:
        with tempfile.TemporaryDirectory() as tmpdir:
            md_path = Path(tmpdir) / "manuscript.md"
            md_path.write_text(markdown_text, encoding="utf-8")

            cmd = [
                "pandoc",
                str(md_path),
                "-o", str(output_path),
                f"--reference-doc={reference_doc}",
            ]
            if use_citeproc:
                cmd.append("--citeproc")

            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                return output_path
            # pandoc failed — fall through to python-docx fallback

    # ── Fallback ──────────────────────────────────────────────────────────────
    return _markdown_to_docx_fallback(markdown_text, output_path, journal_profile)