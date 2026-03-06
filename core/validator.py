from docx import Document
import re

def validate_transformation(original_markdown: str, docx_path: str):
    """
    Performs a sanity check to ensure no significant content was lost
    during the Markdown to DOCX transformation.

    Compares word count and heading count.

    Args:
        original_markdown: The original Markdown text.
        docx_path: Path to the generated .docx file.

    Returns:
        A dictionary with the validation results.
    """
    # 1. Get stats from original Markdown
    original_word_count = len(original_markdown.split())
    original_heading_count = len(re.findall(r"^#", original_markdown, re.MULTILINE))

    # 2. Get stats from generated DOCX
    try:
        document = Document(docx_path)
        docx_text = "\n".join([para.text for para in document.paragraphs])
        docx_word_count = len(docx_text.split())
        
        # Count headings (paragraphs with a style like 'Heading 1', 'Heading 2', etc.)
        docx_heading_count = 0
        for para in document.paragraphs:
            if para.style.name.startswith('Heading'):
                docx_heading_count += 1

    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to read or parse the generated DOCX file: {e}"
        }

    # 3. Compare and generate report
    word_diff = docx_word_count - original_word_count
    heading_diff = docx_heading_count - original_heading_count

    # A simple heuristic for success: word count should not deviate by more than 20%
    # and heading count should be the same or greater.
    is_sane = (abs(word_diff) / original_word_count < 0.2) and (heading_diff >= 0)

    return {
        "status": "completed",
        "is_sane": is_sane,
        "original_word_count": original_word_count,
        "docx_word_count": docx_word_count,
        "word_difference": word_diff,
        "original_heading_count": original_heading_count,
        "docx_heading_count": docx_heading_count,
        "heading_difference": heading_diff,
        "suggestion": "Check the generated document for content loss." if not is_sane else "Transformation seems successful."
    }
